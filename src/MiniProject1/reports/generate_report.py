from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set Times New Roman as default
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_heading('Mini Project 1: Exploratory Data Analysis (EDA)', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Introduction
doc.add_heading('Introduction', level=2)
doc.add_paragraph(
    "This project focuses on performing Exploratory Data Analysis (EDA) "
    "on a customer analytics dataset. The objective was to inspect, clean, "
    "visualize, and extract meaningful insights from the data."
)

# Project Structure
doc.add_heading('Project Structure', level=2)
doc.add_paragraph(
    "The project contains the following components:\n"
    "- customer_analytics.csv (Original Dataset)\n"
    "- customer_analytics_cleaned.csv (Cleaned Dataset)\n"
    "- MiniProject1_EDA.ipynb (Notebook with analysis)\n"
    "- Documentation Report\n"
    "- requirements.txt"
)

# Data Cleaning
doc.add_heading('Data Cleaning', level=2)
doc.add_paragraph(
    "Missing numerical values were filled using median imputation. "
    "Duplicate rows were removed to ensure data integrity."
)

# Visualization
doc.add_heading('Data Visualization', level=2)
doc.add_paragraph(
    "Univariate analysis was performed using histograms and boxplots. "
    "Bivariate analysis was conducted using scatter plots. "
    "A correlation heatmap was generated to identify relationships between variables."
)

# Key Insights
doc.add_heading('Key Insights', level=2)
doc.add_paragraph(
    "1. Moderate correlations were observed between certain numerical variables.\n"
    "2. Outliers were detected in selected features.\n"
    "3. Proper handling of missing data improved dataset reliability."
)

# Conclusion
doc.add_heading('Conclusion', level=2)
doc.add_paragraph(
    "This project strengthened understanding of data preprocessing, "
    "visualization techniques, and structured analytical reporting."
)

# Save document
doc.save("MiniProject1_EDA_Report.docx")

print("Word report generated successfully!")